"""
Export PRISMA 2020 flow diagram + DOCX report dari hasil screening SLR.

Generates:
  - PRISMA 2020 flow diagram (PNG, publication-ready)
  - DOCX report berisi diagram + 27-item PRISMA checklist + abstract checklist
    (semua box di DOCX bisa di-edit manual sebelum submit ke jurnal)

Usage:
  python export_prisma.py
  python export_prisma.py -i "data/Ready (1).ris" -l output/screening_log.json
  python export_prisma.py --registers 0 --other-removed 0 \\
      --not-retrieved 0 --fulltext-assessed 83 \\
      --fulltext-excluded 0 --fulltext-reasons "Wrong topic:0; No full text:0"

Reference: https://www.prisma-statement.org/prisma-2020
"""

import io
import json
import argparse
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── RIS Parser & Duplicate Detection (mirror dari slr_screening.py) ────────────
def parse_ris(filepath):
    articles, current = [], {}
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.rstrip("\n")
            if line.startswith("ER  -"):
                if current:
                    articles.append(current)
                    current = {}
            elif "  - " in line:
                tag = line[:2].strip()
                value = line[6:]
                if tag == "TY":
                    current = {"TY": value}
                elif tag in current:
                    if isinstance(current[tag], list):
                        current[tag].append(value)
                    else:
                        current[tag] = [current[tag], value]
                else:
                    current[tag] = value
    return articles


def count_duplicates(articles):
    seen_titles, seen_dois = {}, {}
    dup = 0
    for i, a in enumerate(articles):
        title = (a.get("TI") or "").lower().strip()
        doi = (a.get("DO") or "").lower().strip()
        if doi and doi in seen_dois:
            dup += 1
        elif title and title in seen_titles:
            dup += 1
        else:
            if doi:
                seen_dois[doi] = i
            if title:
                seen_titles[title] = i
    return dup


# ── Hitung counts buat PRISMA ──────────────────────────────────────────────────
def compute_counts(articles, log, args):
    total = len(articles)
    duplicates = count_duplicates(articles)

    decisions = [v.get("decision", "").lower() for v in log.values()]
    n_included_abstract = sum(1 for d in decisions if d == "include")
    n_excluded_abstract = sum(1 for d in decisions if d == "exclude")

    fulltext_assessed = args.fulltext_assessed if args.fulltext_assessed is not None else n_included_abstract
    fulltext_excluded = args.fulltext_excluded or 0
    not_retrieved = args.not_retrieved or 0

    fulltext_reasons = {}
    if args.fulltext_reasons:
        for pair in args.fulltext_reasons.split(";"):
            pair = pair.strip()
            if not pair:
                continue
            if ":" in pair:
                k, v = pair.rsplit(":", 1)
                try:
                    fulltext_reasons[k.strip()] = int(v.strip())
                except ValueError:
                    pass

    final_included = max(fulltext_assessed - fulltext_excluded, 0)

    return {
        "total_databases": total,
        "total_registers": args.registers or 0,
        "duplicates": duplicates,
        "other_removed": args.other_removed or 0,
        "screened": total - duplicates - (args.other_removed or 0),
        "excluded_abstract": n_excluded_abstract,
        "sought_retrieval": n_included_abstract,
        "not_retrieved": not_retrieved,
        "fulltext_assessed": fulltext_assessed,
        "fulltext_excluded": fulltext_excluded,
        "fulltext_reasons": fulltext_reasons,
        "final_included": final_included,
    }


# ── PRISMA Flow Diagram (matplotlib) ───────────────────────────────────────────
COLORS = {
    "identification": "#d6e4f0",
    "screening": "#fce4d6",
    "included": "#d5e8d4",
    "excluded": "#f8cecc",
    "header": "#2c3e50",
    "text": "#1a1a1a",
    "edge": "#34495e",
}


def _box(ax, x, y, w, h, text, facecolor, fontsize=9, fontweight="normal"):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.05",
        linewidth=1.2, edgecolor=COLORS["edge"], facecolor=facecolor,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center",
        fontsize=fontsize, fontweight=fontweight, color=COLORS["text"],
        wrap=True,
    )


def _section_header(ax, x, y, w, h, text):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.01,rounding_size=0.03",
        linewidth=0, facecolor=COLORS["header"],
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center",
        fontsize=11, fontweight="bold", color="white",
    )


def _arrow(ax, x1, y1, x2, y2):
    arrow = FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle="-|>", mutation_scale=15,
        linewidth=1.3, color=COLORS["edge"],
    )
    ax.add_patch(arrow)


def render_flow_png(counts, output_path, title="PRISMA 2020 Flow Diagram"):
    fig, ax = plt.subplots(figsize=(11, 13))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 14)
    ax.axis("off")

    ax.text(5, 13.5, title, ha="center", va="center",
            fontsize=14, fontweight="bold", color=COLORS["header"])

    # Section labels (kiri)
    _section_header(ax, 0.1, 11.5, 0.55, 1.6, "I\nd\ne\nn\nt.")
    _section_header(ax, 0.1, 6.7, 0.55, 4.5, "S\nc\nr\ne\ne\nn\ni\nn\ng")
    _section_header(ax, 0.1, 1.5, 0.55, 1.6, "I\nn\nc\nl.")

    # IDENTIFICATION
    id_text = (
        f"Records identified from:\n"
        f"  Databases (n = {counts['total_databases']:,})\n"
        f"  Registers (n = {counts['total_registers']:,})"
    )
    _box(ax, 0.9, 11.5, 4.2, 1.6, id_text, COLORS["identification"], fontweight="bold")

    id_rem = (
        f"Records removed before\nscreening:\n"
        f"  Duplicate records (n = {counts['duplicates']:,})\n"
        f"  Other reasons (n = {counts['other_removed']:,})"
    )
    _box(ax, 5.5, 11.5, 4.2, 1.6, id_rem, COLORS["excluded"])

    # arrow id -> id removal (horizontal)
    _arrow(ax, 5.1, 12.3, 5.5, 12.3)

    # SCREENING - row 1: screened + excluded abstract
    scr_text = f"Records screened\n(n = {counts['screened']:,})"
    _box(ax, 0.9, 9.7, 4.2, 1.3, scr_text, COLORS["screening"], fontweight="bold")

    exc_abs = f"Records excluded by\ntitle/abstract screening\n(n = {counts['excluded_abstract']:,})"
    _box(ax, 5.5, 9.7, 4.2, 1.3, exc_abs, COLORS["excluded"])

    _arrow(ax, 3.0, 11.5, 3.0, 11.0)  # id -> screened
    _arrow(ax, 5.1, 10.35, 5.5, 10.35)

    # SCREENING - row 2: sought retrieval + not retrieved
    sought = f"Reports sought for retrieval\n(n = {counts['sought_retrieval']:,})"
    _box(ax, 0.9, 8.0, 4.2, 1.3, sought, COLORS["screening"])

    not_ret = f"Reports not retrieved\n(n = {counts['not_retrieved']:,})"
    _box(ax, 5.5, 8.0, 4.2, 1.3, not_ret, COLORS["excluded"])

    _arrow(ax, 3.0, 9.7, 3.0, 9.3)
    _arrow(ax, 5.1, 8.65, 5.5, 8.65)

    # SCREENING - row 3: full-text assessed + excluded with reasons
    assessed = f"Reports assessed for eligibility\n(n = {counts['fulltext_assessed']:,})"
    _box(ax, 0.9, 6.3, 4.2, 1.3, assessed, COLORS["screening"])

    if counts["fulltext_reasons"]:
        reasons_lines = "\n".join(f"  {k} (n = {v:,})" for k, v in counts["fulltext_reasons"].items())
    else:
        reasons_lines = f"  (specify reasons)"
    exc_ft = f"Reports excluded (n = {counts['fulltext_excluded']:,}):\n{reasons_lines}"
    _box(ax, 5.5, 6.0, 4.2, 1.9, exc_ft, COLORS["excluded"])

    _arrow(ax, 3.0, 8.0, 3.0, 7.6)
    _arrow(ax, 5.1, 6.95, 5.5, 6.95)

    # INCLUDED
    included = (
        f"Studies included in review\n"
        f"(n = {counts['final_included']:,})"
    )
    _box(ax, 0.9, 1.7, 4.2, 1.6, included, COLORS["included"],
         fontsize=11, fontweight="bold")

    _arrow(ax, 3.0, 6.3, 3.0, 3.3)

    # Footer note
    ax.text(
        5, 0.6,
        "Generated by SLR Autopilot · PRISMA 2020 (Page et al., BMJ 2021)",
        ha="center", va="center", fontsize=8, style="italic", color="#666",
    )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(output_path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ── DOCX Report ────────────────────────────────────────────────────────────────
PRISMA_CHECKLIST_2020 = [
    ("TITLE", "1", "Identify the report as a systematic review."),
    ("ABSTRACT", "2", "See the PRISMA 2020 for Abstracts checklist."),
    ("INTRODUCTION", "3", "Rationale: Describe the rationale for the review in the context of existing knowledge."),
    ("INTRODUCTION", "4", "Objectives: Provide an explicit statement of the objective(s) or question(s) the review addresses."),
    ("METHODS", "5", "Eligibility criteria: Specify the inclusion and exclusion criteria for the review and how studies were grouped for the syntheses."),
    ("METHODS", "6", "Information sources: Specify all databases, registers, websites, organisations, reference lists and other sources searched or consulted to identify studies. Specify the date when each source was last searched or consulted."),
    ("METHODS", "7", "Search strategy: Present the full search strategies for all databases, registers and websites, including any filters and limits used."),
    ("METHODS", "8", "Selection process: Specify the methods used to decide whether a study met the inclusion criteria of the review, including how many reviewers screened each record and each report retrieved, whether they worked independently, and if applicable, details of automation tools used in the process."),
    ("METHODS", "9", "Data collection process: Specify the methods used to collect data from reports, including how many reviewers collected data from each report, whether they worked independently, any processes for obtaining or confirming data from study investigators, and if applicable, details of automation tools used in the process."),
    ("METHODS", "10a", "Data items: List and define all outcomes for which data were sought. Specify whether all results that were compatible with each outcome domain in each study were sought."),
    ("METHODS", "10b", "Data items: List and define all other variables for which data were sought (e.g. participant and intervention characteristics, funding sources)."),
    ("METHODS", "11", "Study risk of bias assessment: Specify the methods used to assess risk of bias in the included studies, including details of the tool(s) used, how many reviewers assessed each study and whether they worked independently, and if applicable, details of automation tools used in the process."),
    ("METHODS", "12", "Effect measures: Specify for each outcome the effect measure(s) (e.g. risk ratio, mean difference) used in the synthesis or presentation of results."),
    ("METHODS", "13a", "Synthesis methods: Describe the processes used to decide which studies were eligible for each synthesis."),
    ("METHODS", "13b", "Synthesis methods: Describe any methods required to prepare the data for presentation or synthesis."),
    ("METHODS", "13c", "Synthesis methods: Describe any methods used to tabulate or visually display results of individual studies and syntheses."),
    ("METHODS", "13d", "Synthesis methods: Describe any methods used to synthesize results."),
    ("METHODS", "13e", "Synthesis methods: Describe any methods used to explore possible causes of heterogeneity among study results."),
    ("METHODS", "13f", "Synthesis methods: Describe any sensitivity analyses conducted to assess robustness of the synthesized results."),
    ("METHODS", "14", "Reporting bias assessment: Describe any methods used to assess risk of bias due to missing results in a synthesis."),
    ("METHODS", "15", "Certainty assessment: Describe any methods used to assess certainty (or confidence) in the body of evidence for an outcome."),
    ("RESULTS", "16a", "Study selection: Describe the results of the search and selection process, from the number of records identified in the search to the number of studies included in the review, ideally using a flow diagram."),
    ("RESULTS", "16b", "Study selection: Cite studies that might appear to meet the inclusion criteria, but which were excluded, and explain why they were excluded."),
    ("RESULTS", "17", "Study characteristics: Cite each included study and present its characteristics."),
    ("RESULTS", "18", "Risk of bias in studies: Present assessments of risk of bias for each included study."),
    ("RESULTS", "19", "Results of individual studies: For all outcomes, present, for each study: summary statistics for each group and an effect estimate and its precision."),
    ("RESULTS", "20a", "Results of syntheses: For each synthesis, briefly summarise the characteristics and risk of bias among contributing studies."),
    ("RESULTS", "20b", "Results of syntheses: Present results of all statistical syntheses conducted."),
    ("RESULTS", "20c", "Results of syntheses: Present results of all investigations of possible causes of heterogeneity among study results."),
    ("RESULTS", "20d", "Results of syntheses: Present results of all sensitivity analyses conducted."),
    ("RESULTS", "21", "Reporting biases: Present assessments of risk of bias due to missing results for each synthesis."),
    ("RESULTS", "22", "Certainty of evidence: Present assessments of certainty (or confidence) in the body of evidence for each outcome assessed."),
    ("DISCUSSION", "23a", "Discussion: Provide a general interpretation of the results in the context of other evidence."),
    ("DISCUSSION", "23b", "Discussion: Discuss any limitations of the evidence included in the review."),
    ("DISCUSSION", "23c", "Discussion: Discuss any limitations of the review processes used."),
    ("DISCUSSION", "23d", "Discussion: Discuss implications of the results for practice, policy, and future research."),
    ("OTHER", "24a", "Registration and protocol: Provide registration information for the review, including register name and registration number, or state that the review was not registered."),
    ("OTHER", "24b", "Registration and protocol: Indicate where the review protocol can be accessed, or state that a protocol was not prepared."),
    ("OTHER", "24c", "Registration and protocol: Describe and explain any amendments to information provided at registration or in the protocol."),
    ("OTHER", "25", "Support: Describe sources of financial or non-financial support for the review, and the role of the funders or sponsors in the review."),
    ("OTHER", "26", "Competing interests: Declare any competing interests of review authors."),
    ("OTHER", "27", "Availability of data, code and other materials: Report which of the following are publicly available and where they can be found: template data collection forms; data extracted from included studies; data used for all analyses; analytic code; any other materials used in the review."),
]


PRISMA_ABSTRACT_CHECKLIST = [
    ("TITLE", "1", "Identify the report as a systematic review."),
    ("BACKGROUND", "2", "Objectives: Provide an explicit statement of the main objective(s) or question(s) the review addresses."),
    ("METHODS", "3", "Eligibility criteria: Specify the inclusion and exclusion criteria for the review."),
    ("METHODS", "4", "Information sources: Specify the information sources (e.g. databases, registers) used to identify studies and the date when each was last searched."),
    ("METHODS", "5", "Risk of bias: Specify the methods used to assess risk of bias in the included studies."),
    ("METHODS", "6", "Synthesis of results: Specify the methods used to present and synthesise results."),
    ("RESULTS", "7", "Included studies: Give the total number of included studies and participants and summarise relevant characteristics of studies."),
    ("RESULTS", "8", "Synthesis of results: Present results for main outcomes, preferably indicating the number of included studies and participants for each. If meta-analysis was done, report the summary estimate and confidence/credible interval."),
    ("DISCUSSION", "9", "Limitations of evidence: Provide a brief summary of the limitations of the evidence included in the review (e.g. study risk of bias, inconsistency and imprecision)."),
    ("DISCUSSION", "10", "Interpretation: Provide a general interpretation of the results and important implications."),
    ("OTHER", "11", "Funding: Specify the primary source of funding for the review."),
    ("OTHER", "12", "Registration: Provide the register name and registration number."),
]


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "808080")
        borders.append(b)
    tc_pr.append(borders)


def _heading_para(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return h


def _add_flow_table(doc, counts):
    """Editable table-based PRISMA flow — author can tweak text directly in Word."""
    rows_def = [
        ("IDENTIFICATION", COLORS["identification"], [
            (f"Records identified from:\n  Databases (n = {counts['total_databases']:,})\n  Registers (n = {counts['total_registers']:,})",
             f"Records removed before screening:\n  Duplicates (n = {counts['duplicates']:,})\n  Other reasons (n = {counts['other_removed']:,})"),
        ]),
        ("SCREENING", COLORS["screening"], [
            (f"Records screened (n = {counts['screened']:,})",
             f"Records excluded by title/abstract (n = {counts['excluded_abstract']:,})"),
            (f"Reports sought for retrieval (n = {counts['sought_retrieval']:,})",
             f"Reports not retrieved (n = {counts['not_retrieved']:,})"),
            (f"Reports assessed for eligibility (n = {counts['fulltext_assessed']:,})",
             _format_fulltext_reasons(counts)),
        ]),
        ("INCLUDED", COLORS["included"], [
            (f"Studies included in review (n = {counts['final_included']:,})", ""),
        ]),
    ]

    for section, color, rows in rows_def:
        sec_table = doc.add_table(rows=1, cols=1)
        sec_table.autofit = False
        sec_cell = sec_table.rows[0].cells[0]
        _shade_cell(sec_cell, COLORS["header"])
        _set_cell_borders(sec_cell)
        p = sec_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(section)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)

        body = doc.add_table(rows=len(rows), cols=2)
        body.autofit = False
        for i, (left, right) in enumerate(rows):
            lc, rc = body.rows[i].cells
            for cell, text, shade in ((lc, left, color), (rc, right, COLORS["excluded"] if right else "FFFFFF")):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if text:
                    _shade_cell(cell, shade)
                _set_cell_borders(cell)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                run.font.size = Pt(9)
        doc.add_paragraph()


def _format_fulltext_reasons(counts):
    n_exc = counts["fulltext_excluded"]
    if counts["fulltext_reasons"]:
        lines = "\n".join(f"  {k} (n = {v:,})" for k, v in counts["fulltext_reasons"].items())
    else:
        lines = "  (specify reasons)"
    return f"Reports excluded (n = {n_exc:,}):\n{lines}"


def _add_checklist_table(doc, checklist, title):
    _heading_para(doc, title, level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    headers = ["Section", "Item #", "Checklist item", "Reported on page #"]
    widths = [Cm(2.5), Cm(1.5), Cm(10.0), Cm(3.0)]
    for cell, text, w in zip(hdr, headers, widths):
        _shade_cell(cell, COLORS["header"])
        _set_cell_borders(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.width = w

    for section, num, desc in checklist:
        row = table.add_row().cells
        row[0].text = section
        row[1].text = num
        row[2].text = desc
        row[3].text = ""
        for cell, w in zip(row, widths):
            _set_cell_borders(cell)
            cell.width = w
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def render_docx(counts, flow_png_path, output_path, title, args):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title block
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"PRISMA 2020 Report  ·  Generated by SLR Autopilot\n"
        f"Input: {Path(args.input).name}  ·  Log: {Path(args.log).name}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Section 1: embedded diagram
    _heading_para(doc, "1. PRISMA 2020 Flow Diagram", level=1)
    doc.add_paragraph(
        "Diagram di bawah ini di-generate otomatis dari hasil screening. "
        "Versi editable (table-based) ada di Section 2."
    ).runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(flow_png_path), width=Inches(6.0))

    doc.add_page_break()

    # Section 2: editable table version
    _heading_para(doc, "2. Flow Diagram (Editable)", level=1)
    doc.add_paragraph(
        "Edit teks di box manapun langsung di Word — angka diisi dari hasil screening, "
        "tapi bisa disesuaikan sebelum submit ke jurnal."
    ).runs[0].font.size = Pt(10)
    _add_flow_table(doc, counts)

    doc.add_page_break()

    # Section 3: 27-item checklist
    _heading_para(doc, "3. PRISMA 2020 Checklist (27 items)", level=1)
    doc.add_paragraph(
        "Isi kolom \"Reported on page #\" sesuai lokasi di manuscript anda."
    ).runs[0].font.size = Pt(10)
    _add_checklist_table(doc, PRISMA_CHECKLIST_2020, "Main Checklist")

    doc.add_page_break()

    # Section 4: Abstract checklist
    _heading_para(doc, "4. PRISMA 2020 for Abstracts (12 items)", level=1)
    _add_checklist_table(doc, PRISMA_ABSTRACT_CHECKLIST, "Abstract Checklist")

    # Section 5: source citation
    doc.add_paragraph()
    _heading_para(doc, "Source", level=2)
    cite = doc.add_paragraph()
    cite.add_run(
        "Page MJ, McKenzie JE, Bossuyt PM, et al. The PRISMA 2020 statement: "
        "an updated guideline for reporting systematic reviews. BMJ 2021;372:n71. "
        "doi:10.1136/bmj.n71  ·  https://www.prisma-statement.org/prisma-2020"
    ).font.size = Pt(9)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ── CLI ────────────────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(
        description="Generate PRISMA 2020 flow diagram (PNG) + DOCX report dari hasil screening SLR",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument("-i", "--input", default="sample.ris", help="Input RIS file")
    p.add_argument("-l", "--log", default="output/screening_log.json", help="Screening log JSON")
    p.add_argument("--png", default="output/prisma_flow.png", help="Output PNG path")
    p.add_argument("--docx", default="output/prisma_report.docx", help="Output DOCX path")
    p.add_argument("--title", default="PRISMA 2020 Flow Diagram", help="Diagram & report title")

    p.add_argument("--registers", type=int, default=0,
                   help="Records identified from registers (default: 0)")
    p.add_argument("--other-removed", type=int, default=0,
                   help="Records removed for reasons other than duplicates (default: 0)")
    p.add_argument("--not-retrieved", type=int, default=0,
                   help="Reports not retrieved (default: 0)")
    p.add_argument("--fulltext-assessed", type=int, default=None,
                   help="Reports assessed for eligibility (default: jumlah include dari abstract screening)")
    p.add_argument("--fulltext-excluded", type=int, default=0,
                   help="Reports excluded at full-text stage (default: 0)")
    p.add_argument("--fulltext-reasons", type=str, default="",
                   help='Reasons for full-text exclusion, format "Reason:N; Reason:N"')
    return p.parse_args()


def main():
    args = parse_args()

    print(f"Membaca RIS: {args.input}")
    articles = parse_ris(args.input)

    print(f"Membaca log screening: {args.log}")
    with open(args.log, "r") as f:
        log = json.load(f)

    counts = compute_counts(articles, log, args)

    print("Counts:")
    for k, v in counts.items():
        print(f"  {k:24s} = {v}")

    print(f"\nMembuat PRISMA flow diagram (PNG)...")
    render_flow_png(counts, args.png, title=args.title)
    print(f"  -> {args.png}")

    print(f"Membuat DOCX report...")
    render_docx(counts, args.png, args.docx, args.title, args)
    print(f"  -> {args.docx}")

    print(f"\nSelesai!")
    print(f"  Studies included in review: {counts['final_included']:,}")


if __name__ == "__main__":
    main()
